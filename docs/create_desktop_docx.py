#!/usr/bin/env python3
"""
Create a separate DOCX document for Desktop Application screenshots,
following the same format as the web application screenshots section.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCREENSHOT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots')
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Tampilan_Layar_Aplikasi_Desktop.docx')


def set_paragraph_spacing(para, before=0, after=0):
    """Set paragraph spacing in points."""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = pPr.makeelement(qn('w:spacing'), {})
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before * 20)))  # twips
    spacing.set(qn('w:after'), str(int(after * 20)))


def add_image_with_caption(doc, filepath, caption, img_width=5.5):
    """Add an image centered with caption below, matching web app format."""
    if not os.path.exists(filepath):
        print(f"  WARNING: {filepath} not found, skipping")
        return False

    # Image paragraph (centered)
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(filepath, width=Inches(img_width))
    set_paragraph_spacing(img_para, before=6, after=3)

    # Caption paragraph (centered, normal style)
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.font.size = Pt(10)
    set_paragraph_spacing(cap_para, before=3, after=12)

    return True


def main():
    print("=" * 60)
    print("Creating Desktop Application Screenshots Document")
    print("=" * 60)

    if not os.path.exists(SCREENSHOT_DIR):
        print(f"Error: {SCREENSHOT_DIR} not found.")
        return

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # =========================================================
    # TITLE
    # =========================================================
    title = doc.add_heading('Tampilan Layar Aplikasi Presensi Sholat Berbasis Desktop', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_run = intro.add_run(
        'Bagian ini menampilkan antarmuka pengguna (UI) dari aplikasi presensi sholat '
        'berbasis Desktop yang telah dikembangkan menggunakan Electron, React, dan Tailwind CSS. '
        'Setiap halaman ditampilkan untuk masing-masing peran pengguna: Administrator, Wali Kelas, '
        'Guru, dan Siswa.'
    )
    intro_run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # =========================================================
    # 6.1 AUTENTIKASI
    # =========================================================
    doc.add_heading('6.1 Tampilan Autentikasi', level=2)

    auth_intro = doc.add_paragraph()
    auth_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    auth_intro.add_run(
        'Saat membuka aplikasi berbasis desktop, pengguna pertama kali dihadapkan pada halaman '
        '"Masuk" yang menampilkan kolom input untuk NIS/NIP dan Kata Sandi, disertai tombol "Masuk", '
        'tautan "Lupa Kata Sandi?", dan "Belum punya akun? Daftar".'
    )

    screenshots_auth = [
        ("Gambar 6.0 Halaman Login.png", "Gambar 6.1 Antarmuka Autentikasi - Login"),
    ]
    for filename, caption in screenshots_auth:
        add_image_with_caption(doc, os.path.join(SCREENSHOT_DIR, filename), caption)

    # =========================================================
    # 6.2 PORTAL ADMIN
    # =========================================================
    doc.add_heading('6.2 Portal Administrator', level=2)

    admin_intro = doc.add_paragraph()
    admin_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    admin_intro.add_run(
        'Sebagai admin, tampilan beranda menampilkan ringkasan data kehadiran siswa, '
        'jadwal sholat terdekat, dan grafik tren kehadiran. Melalui menu navigasi sidebar, '
        'admin dapat mengakses seluruh fitur pengelolaan sistem presensi sholat.'
    )

    screenshots_admin = [
        ("Gambar 6.1 Halaman Beranda Admin.png", "Gambar 6.2 Antarmuka Admin - Dashboard"),
        ("Gambar 6.2 Halaman Jadwal Admin.png", "Gambar 6.3 Antarmuka Admin - Jadwal"),
        ("Gambar 6.3 Halaman Kelola Siswa Admin.png", "Gambar 6.4 Antarmuka Admin - Kelola Siswa"),
        ("Gambar 6.4 Halaman Kelola Kelas Admin.png", "Gambar 6.5 Antarmuka Admin - Kelola Kelas"),
        ("Gambar 6.5 Halaman Kelola Guru Admin.png", "Gambar 6.6 Antarmuka Admin - Kelola Guru"),
        ("Gambar 6.6 Halaman Presensi Admin.png", "Gambar 6.7 Antarmuka Admin - Presensi"),
        ("Gambar 6.7 Halaman Pengajuan Izin Admin.png", "Gambar 6.8 Antarmuka Admin - Pengajuan Izin"),
        ("Gambar 6.8 Halaman Laporan Admin.png", "Gambar 6.9 Antarmuka Admin - Laporan"),
        ("Gambar 6.9 Halaman QR Code Admin.png", "Gambar 6.10 Antarmuka Admin - QR Code"),
        ("Gambar 6.10 Halaman Siswa Belum Terdaftar Admin.png", "Gambar 6.11 Antarmuka Admin - Siswa Belum Terdaftar"),
        ("Gambar 6.11 Halaman Perangkat Siswa Admin.png", "Gambar 6.12 Antarmuka Admin - Perangkat Siswa"),
        ("Gambar 6.12 Halaman Profile Admin.png", "Gambar 6.13 Antarmuka Admin - Akun"),
        ("Gambar 6.13 Halaman Pengaturan Admin.png", "Gambar 6.14 Antarmuka Admin - Pengaturan"),
    ]
    for filename, caption in screenshots_admin:
        add_image_with_caption(doc, os.path.join(SCREENSHOT_DIR, filename), caption)

    # =========================================================
    # 6.3 PORTAL WALI KELAS
    # =========================================================
    doc.add_heading('6.3 Portal Wali Kelas', level=2)

    wali_intro = doc.add_paragraph()
    wali_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    wali_intro.add_run(
        'Berfungsi sebagai dashboard utama yang menyajikan ringkasan statistik presensi siswa '
        'kelas yang diampu secara real-time. Wali kelas dapat melihat rekapitulasi harian, '
        'jadwal sholat, presensi siswa, verifikasi pengajuan izin, serta analisis laporan kehadiran.'
    )

    screenshots_wali = [
        ("Gambar 6.14 Halaman Beranda Guru.png", "Gambar 6.15 Antarmuka Wali Kelas - Dashboard"),
        ("Gambar 6.15 Halaman Jadwal Guru.png", "Gambar 6.16 Antarmuka Wali Kelas - Jadwal"),
        ("Gambar 6.16 Halaman Presensi Guru.png", "Gambar 6.17 Antarmuka Wali Kelas - Presensi"),
        ("Gambar 6.17 Halaman Pengajuan Izin Guru.png", "Gambar 6.18 Antarmuka Wali Kelas - Pengajuan Izin"),
        ("Gambar 6.18 Halaman Laporan Guru.png", "Gambar 6.19 Antarmuka Wali Kelas - Laporan"),
        ("Gambar 6.19 Halaman Siswa Belum Terdaftar Guru.png", "Gambar 6.20 Antarmuka Wali Kelas - Siswa Belum Terdaftar"),
        ("Gambar 6.20 Halaman Profile Guru.png", "Gambar 6.21 Antarmuka Wali Kelas - Akun"),
        ("Gambar 6.21 Halaman Pengaturan Guru.png", "Gambar 6.22 Antarmuka Wali Kelas - Pengaturan"),
    ]
    for filename, caption in screenshots_wali:
        add_image_with_caption(doc, os.path.join(SCREENSHOT_DIR, filename), caption)

    # =========================================================
    # 6.4 PORTAL GURU
    # =========================================================
    doc.add_heading('6.4 Portal Guru', level=2)

    guru_intro = doc.add_paragraph()
    guru_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    guru_intro.add_run(
        'Guru dapat melihat ringkasan statistik presensi siswa secara real-time, '
        'jadwal sholat, rekapitulasi presensi, pengajuan izin, serta laporan kehadiran. '
        'Tampilan dan fitur yang tersedia serupa dengan portal wali kelas.'
    )

    screenshots_guru = [
        ("Gambar 6.14 Halaman Beranda Guru.png", "Gambar 6.23 Antarmuka Guru - Dashboard"),
        ("Gambar 6.15 Halaman Jadwal Guru.png", "Gambar 6.24 Antarmuka Guru - Jadwal"),
        ("Gambar 6.16 Halaman Presensi Guru.png", "Gambar 6.25 Antarmuka Guru - Presensi"),
        ("Gambar 6.17 Halaman Pengajuan Izin Guru.png", "Gambar 6.26 Antarmuka Guru - Pengajuan Izin"),
        ("Gambar 6.18 Halaman Laporan Guru.png", "Gambar 6.27 Antarmuka Guru - Laporan"),
        ("Gambar 6.19 Halaman Siswa Belum Terdaftar Guru.png", "Gambar 6.28 Antarmuka Guru - Siswa Belum Terdaftar"),
        ("Gambar 6.20 Halaman Profile Guru.png", "Gambar 6.29 Antarmuka Guru - Akun"),
        ("Gambar 6.21 Halaman Pengaturan Guru.png", "Gambar 6.30 Antarmuka Guru - Pengaturan"),
    ]
    for filename, caption in screenshots_guru:
        add_image_with_caption(doc, os.path.join(SCREENSHOT_DIR, filename), caption)

    # =========================================================
    # 6.5 PORTAL SISWA
    # =========================================================
    doc.add_heading('6.5 Portal Siswa', level=2)

    siswa_intro = doc.add_paragraph()
    siswa_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    siswa_intro.add_run(
        'Siswa diarahkan ke halaman beranda yang menampilkan statistik kehadiran pribadi, '
        'jadwal sholat hari ini, dan riwayat absensi. Siswa dapat melakukan presensi melalui '
        'pindai QR Code atau kode manual, mengajukan izin/sakit, serta mengelola profil akun.'
    )

    screenshots_siswa = [
        ("Gambar 6.22 Halaman Beranda Siswa.png", "Gambar 6.31 Antarmuka Siswa - Dashboard"),
        ("Gambar 6.23 Halaman Pindai QR Siswa.png", "Gambar 6.32 Antarmuka Siswa - Pindai QR"),
        ("Gambar 6.24 Halaman Izin Siswa.png", "Gambar 6.33 Antarmuka Siswa - Pengajuan Izin"),
        ("Gambar 6.25 Halaman Profil Siswa.png", "Gambar 6.34 Antarmuka Siswa - Akun"),
        ("Gambar 6.26 Halaman Pengaturan Siswa.png", "Gambar 6.35 Antarmuka Siswa - Pengaturan"),
    ]
    for filename, caption in screenshots_siswa:
        add_image_with_caption(doc, os.path.join(SCREENSHOT_DIR, filename), caption)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\n{'=' * 60}")
    print(f"Document saved to: {OUTPUT_PATH}")
    print("=" * 60)


if __name__ == '__main__':
    main()
