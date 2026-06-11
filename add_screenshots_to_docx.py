#!/usr/bin/env python3
"""
Append desktop app screenshots to the Laporan DOCX following the same format
as the web app screenshots (section 5.1).
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCREENSHOT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots')
DOCX_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Laporan_TAXIRPL2SEMSTWO111_Updated.docx')

# Format matches web app section: "Gambar X.X Antarmuka [Role] - [Page Name]"
SCREENSHOT_DATA = [
    ("6. Tampilan Layar Aplikasi Desktop", [
        ("Gambar 6.0 Halaman Login.png", "Gambar 6.0 Antarmuka Autentikasi - Login"),
    ]),
    ("6.1 Portal Administrator", [
        ("Gambar 6.1 Halaman Beranda Admin.png", "Gambar 6.1 Antarmuka Admin - Dashboard"),
        ("Gambar 6.2 Halaman Jadwal Admin.png", "Gambar 6.2 Antarmuka Admin - Jadwal"),
        ("Gambar 6.3 Halaman Kelola Siswa Admin.png", "Gambar 6.3 Antarmuka Admin - Kelola Siswa"),
        ("Gambar 6.4 Halaman Kelola Kelas Admin.png", "Gambar 6.4 Antarmuka Admin - Kelola Kelas"),
        ("Gambar 6.5 Halaman Kelola Guru Admin.png", "Gambar 6.5 Antarmuka Admin - Kelola Guru"),
        ("Gambar 6.6 Halaman Presensi Admin.png", "Gambar 6.6 Antarmuka Admin - Presensi"),
        ("Gambar 6.7 Halaman Pengajuan Izin Admin.png", "Gambar 6.7 Antarmuka Admin - Pengajuan Izin"),
        ("Gambar 6.8 Halaman Laporan Admin.png", "Gambar 6.8 Antarmuka Admin - Laporan"),
        ("Gambar 6.9 Halaman QR Code Admin.png", "Gambar 6.9 Antarmuka Admin - QR Code"),
        ("Gambar 6.10 Halaman Siswa Belum Terdaftar Admin.png", "Gambar 6.10 Antarmuka Admin - Siswa Belum Terdaftar"),
        ("Gambar 6.11 Halaman Perangkat Siswa Admin.png", "Gambar 6.11 Antarmuka Admin - Perangkat Siswa"),
        ("Gambar 6.12 Halaman Profile Admin.png", "Gambar 6.12 Antarmuka Admin - Akun"),
        ("Gambar 6.13 Halaman Pengaturan Admin.png", "Gambar 6.13 Antarmuka Admin - Pengaturan"),
    ]),
    ("6.2 Portal Wali Kelas", [
        ("Gambar 6.14 Halaman Beranda Guru.png", "Gambar 6.14 Antarmuka Wali Kelas - Dashboard"),
        ("Gambar 6.15 Halaman Jadwal Guru.png", "Gambar 6.15 Antarmuka Wali Kelas - Jadwal"),
        ("Gambar 6.16 Halaman Presensi Guru.png", "Gambar 6.16 Antarmuka Wali Kelas - Presensi"),
        ("Gambar 6.17 Halaman Pengajuan Izin Guru.png", "Gambar 6.17 Antarmuka Wali Kelas - Pengajuan Izin"),
        ("Gambar 6.18 Halaman Laporan Guru.png", "Gambar 6.18 Antarmuka Wali Kelas - Laporan"),
        ("Gambar 6.19 Halaman Siswa Belum Terdaftar Guru.png", "Gambar 6.19 Antarmuka Wali Kelas - Siswa Belum Terdaftar"),
        ("Gambar 6.20 Halaman Profile Guru.png", "Gambar 6.20 Antarmuka Wali Kelas - Akun"),
        ("Gambar 6.21 Halaman Pengaturan Guru.png", "Gambar 6.21 Antarmuka Wali Kelas - Pengaturan"),
    ]),
    ("6.3 Portal Siswa", [
        ("Gambar 6.22 Halaman Beranda Siswa.png", "Gambar 6.22 Antarmuka Siswa - Dashboard"),
        ("Gambar 6.23 Halaman Pindai QR Siswa.png", "Gambar 6.23 Antarmuka Siswa - Pindai QR"),
        ("Gambar 6.24 Halaman Izin Siswa.png", "Gambar 6.24 Antarmuka Siswa - Pengajuan Izin"),
        ("Gambar 6.25 Halaman Profil Siswa.png", "Gambar 6.25 Antarmuka Siswa - Akun"),
        ("Gambar 6.26 Halaman Pengaturan Siswa.png", "Gambar 6.26 Antarmuka Siswa - Pengaturan"),
    ]),
]


def remove_existing_chapter6(doc):
    """Remove existing chapter 6 content that was previously added."""
    indices_to_remove = []
    in_chapter6 = False
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith('6. Tampilan Layar Aplikasi'):
            in_chapter6 = True
        if in_chapter6:
            indices_to_remove.append(i)
    
    # Remove in reverse order to preserve indices
    for i in reversed(indices_to_remove):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
    
    return len(indices_to_remove)


def main():
    print("=" * 60)
    print("Appending Desktop App Screenshots to DOCX")
    print("=" * 60)

    if not os.path.exists(SCREENSHOT_DIR):
        print(f"Error: {SCREENSHOT_DIR} not found.")
        return

    if not os.path.exists(DOCX_PATH):
        print(f"Error: {DOCX_PATH} not found.")
        return

    print(f"Loading: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    # Remove any previously added chapter 6
    removed = remove_existing_chapter6(doc)
    if removed:
        print(f"Removed {removed} existing chapter 6 paragraphs")

    total = 0
    for heading_text, screenshots in SCREENSHOT_DATA:
        print(f"\n--- {heading_text} ---")
        
        # Add heading (Heading 1 for main, Heading 3 for portals)
        level = 1 if heading_text.startswith('6.') and '.' not in heading_text[2:] else 3
        if heading_text.startswith('6.') and heading_text[2:].strip()[0].isdigit() and '.' not in heading_text[2:3]:
            level = 1
        else:
            level = 3
        
        heading = doc.add_heading(heading_text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for filename, caption in screenshots:
            filepath = os.path.join(SCREENSHOT_DIR, filename)
            if not os.path.exists(filepath):
                print(f"  WARNING: {filepath} not found, skipping")
                continue

            # Image paragraph (centered, empty text) - matches web format
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(filepath, width=Inches(5.5))

            # Caption paragraph (centered) - matches web format
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            cap_run.font.size = Pt(10)

            print(f"  + {caption}")
            total += 1

    doc.save(DOCX_PATH)
    print(f"\n{'=' * 60}")
    print(f"Saved to: {DOCX_PATH}")
    print(f"Total images added: {total}")
    print("=" * 60)


if __name__ == '__main__':
    main()
